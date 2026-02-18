"""
Answer Key Service
==================
Parses answer keys from multiple file formats (CSV, XLSX, PDF, Images)
and normalizes them into a static AnswerKey object.
"""

import os
import re
import json
import csv
import base64
import requests
from typing import Dict, Optional
from models import AnswerKey, AnswerKeyEntry
from google.oauth2 import service_account
import google.auth.transport.requests
import google.auth

class AnswerKeyService:
    def __init__(self):
        self.project_id = "project-75abf07c-e594-4660-ab7"
        self.location = "us-central1"
        self.model_id = "gemini-2.5-flash-lite"
        self.creds = None
        
        # Try finding Service Account credentials
        creds_path = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
        if creds_path and os.path.exists(creds_path):
            try:
                self.creds = service_account.Credentials.from_service_account_file(
                    creds_path,
                    scopes=["https://www.googleapis.com/auth/cloud-platform"]
                )
                if hasattr(self.creds, "project_id") and self.creds.project_id:
                     self.project_id = self.creds.project_id
            except Exception as e:
                print(f"⚠️ Failed to load Service Account in AnswerKeyService: {e}")

        self.vertex_api_key = os.getenv("VERTEX_AI_API_KEY")

        # Construct URL
        if self.creds:
             self.vertex_url = (
                f"https://{self.location}-aiplatform.googleapis.com/v1/"
                f"projects/{self.project_id}/locations/{self.location}/"
                f"publishers/google/models/{self.model_id}:streamGenerateContent"
             )
        elif self.vertex_api_key:
             self.vertex_url = (
                f"https://aiplatform.googleapis.com/v1/publishers/google/models/"
                f"{self.model_id}:streamGenerateContent?key={self.vertex_api_key}"
             )
        else:
            self.vertex_url = None
            print("⚠️  VERTEX_AI_API_KEY not set and no Service Account found — OCR-based answer key parsing won't work")

    def _get_auth_header(self):
        """Get Authorization header with Bearer token if using Service Account."""
        if self.creds:
            try:
                auth_req = google.auth.transport.requests.Request()
                self.creds.refresh(auth_req)
                return {"Authorization": f"Bearer {self.creds.token}", "Content-Type": "application/json"}
            except Exception as e:
                print(f"❌ Failed to refresh token: {e}")
                return {"Content-Type": "application/json"} 
        return {"Content-Type": "application/json"}

    # ──────────────────────────────────────────
    #  Public API
    # ──────────────────────────────────────────

    def extract_answer_key(self, file_path: str, mime_type: str = None) -> AnswerKey:
        """
        Main entry point. Detects format and routes to the right parser.
        
        Args:
            file_path: Local path to the downloaded answer key file.
            mime_type: Optional MIME type hint from Drive.
            
        Returns:
            AnswerKey object with all answers normalized.
            
        Raises:
            ValueError: If the file format is unsupported or parsing fails.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Answer key file not found: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        print(f"📋 Parsing answer key: {os.path.basename(file_path)} (ext={ext}, mime={mime_type})")

        # Route by extension first, then MIME type as fallback
        if ext == ".csv" or mime_type == "text/csv":
            raw = self._parse_csv(file_path)
        elif ext in (".xlsx", ".xls") or (mime_type and "spreadsheet" in mime_type):
            raw = self._parse_excel(file_path)
        elif ext == ".pdf" or mime_type == "application/pdf":
            raw = self._parse_with_ocr(file_path, mime_type="application/pdf")
        elif ext in (".png", ".jpg", ".jpeg") or (mime_type and "image/" in mime_type):
            raw = self._parse_with_ocr(file_path, mime_type=mime_type)
        elif ext == ".docx":
            raw = self._parse_docx(file_path)
        elif ext == ".txt":
            raw = self._parse_text(file_path)
        else:
            raise ValueError(
                f"Unsupported answer key format: ext={ext}, mime={mime_type}. "
                f"Supported: .csv, .xlsx, .xls, .pdf, .png, .jpg, .jpeg, .docx, .txt"
            )

        # Normalize raw dict → AnswerKey model
        answer_key = self._normalize(raw, source_file=os.path.basename(file_path))
        print(f"✅ Answer key extracted: {answer_key.total_questions} questions")
        
        # PERSISTENCE: Save to disk automatically
        self.save_to_disk(answer_key)
        
        return answer_key

    def save_to_disk(self, answer_key: AnswerKey, filename: str = "current_answer_key.json"):
        """Save the AnswerKey object to a JSON file."""
        try:
            with open(filename, 'w') as f:
                f.write(answer_key.model_dump_json(indent=2))
            print(f"💾 Saved answer key to {filename}")
        except Exception as e:
            print(f"⚠️ Failed to save answer key to disk: {e}")

    def load_from_disk(self, filename: str = "current_answer_key.json") -> Optional[AnswerKey]:
        """Load the AnswerKey object from a JSON file."""
        if not os.path.exists(filename):
            return None
        
        try:
            with open(filename, 'r') as f:
                data = json.load(f)
            
            # Reconstruct Dictionary with integer keys for answers
            answers_data = data.get("answers", {})
            converted_answers = {}
            for k, v in answers_data.items():
                converted_answers[int(k)] = AnswerKeyEntry(**v)
            
            return AnswerKey(
                total_questions=data["total_questions"],
                answers=converted_answers,
                negative_marking=data.get("negative_marking", 0.0),
                metadata=data.get("metadata", {})
            )
        except Exception as e:
            print(f"⚠️ Failed to load answer key from disk: {e}")
            return None

    # ──────────────────────────────────────────
    #  Format-Specific Parsers
    # ──────────────────────────────────────────

    def _parse_csv(self, file_path: str) -> Dict:
        """
        Parse CSV with expected columns: question_number, correct_option, marks (optional)
        
        Also handles simpler formats:
        - Two columns: question_number, correct_option
        - Single column with "Q1: A" style entries
        """
        answers = {}

        with open(file_path, 'r', encoding='utf-8-sig') as f:
            # Sniff the dialect
            sample = f.read(2048)
            f.seek(0)
            
            # Try standard CSV parsing
            try:
                sniffer = csv.Sniffer()
                has_header = sniffer.has_header(sample)
            except csv.Error:
                has_header = True  # assume header

            reader = csv.reader(f)
            headers = None

            if has_header:
                headers = [h.strip().lower() for h in next(reader)]

            # Detect column indices
            q_col, opt_col, marks_col = self._detect_csv_columns(headers)

            for row in reader:
                if not row or all(cell.strip() == '' for cell in row):
                    continue

                try:
                    if q_col is not None and opt_col is not None:
                        q_num = self._parse_question_number(row[q_col].strip())
                        option = row[opt_col].strip().upper()
                        marks = 1.0
                        if marks_col is not None and marks_col < len(row):
                            try:
                                marks = float(row[marks_col].strip())
                            except (ValueError, IndexError):
                                marks = 1.0
                        answers[q_num] = {"correct_option": option, "marks": marks}
                    else:
                        # Try to parse "Q1: A" style from first column
                        parsed = self._parse_inline_answer(row[0])
                        if parsed:
                            answers[parsed[0]] = {"correct_option": parsed[1], "marks": 1.0}
                except (ValueError, IndexError) as e:
                    print(f"  ⚠️  Skipping row {row}: {e}")
                    continue

        if not answers:
            raise ValueError(f"No answers could be parsed from CSV: {file_path}")

        return {"answers": answers}

    def _parse_excel(self, file_path: str) -> Dict:
        """Parse XLSX/XLS file using openpyxl."""
        try:
            import openpyxl
        except ImportError:
            raise ImportError("openpyxl is required for Excel parsing. Run: pip install openpyxl")

        wb = openpyxl.load_workbook(file_path, read_only=True)
        ws = wb.active
        answers = {}

        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            raise ValueError(f"Empty Excel file: {file_path}")

        # Check if first row is header
        first_row = [str(cell).strip().lower() if cell else '' for cell in rows[0]]
        header_keywords = {'question', 'q', 'number', 'option', 'answer', 'correct', 'marks'}
        has_header = any(kw in cell for cell in first_row for kw in header_keywords)

        start_idx = 1 if has_header else 0
        headers = first_row if has_header else None

        q_col, opt_col, marks_col = self._detect_csv_columns(headers)
        # Default column positions if detection fails
        if q_col is None:
            q_col = 0
        if opt_col is None:
            opt_col = 1

        for row in rows[start_idx:]:
            if not row or all(cell is None for cell in row):
                continue
            try:
                q_num = self._parse_question_number(str(row[q_col]).strip())
                option = str(row[opt_col]).strip().upper()
                marks = 1.0
                if marks_col is not None and marks_col < len(row) and row[marks_col]:
                    try:
                        marks = float(row[marks_col])
                    except (ValueError, TypeError):
                        marks = 1.0
                answers[q_num] = {"correct_option": option, "marks": marks}
            except (ValueError, IndexError):
                continue

        wb.close()

        if not answers:
            raise ValueError(f"No answers could be parsed from Excel: {file_path}")

        return {"answers": answers}

    def _parse_docx(self, file_path: str) -> Dict:
        """Parse DOCX by extracting text and using pattern matching."""
        try:
            import docx
        except ImportError:
            # Fallback to OCR-based approach
            print("  python-docx not installed, falling back to OCR")
            return self._parse_with_ocr(file_path, mime_type="application/pdf")

        doc = docx.Document(file_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text for cell in row.cells)
                full_text += "\n" + row_text

        return self._parse_text_content(full_text)

    def _parse_text(self, file_path: str) -> Dict:
        """Parse plain text answer key."""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return self._parse_text_content(content)

    def _parse_text_content(self, text: str) -> Dict:
        """
        Parse free-form text that contains answers like:
        - Q1: A
        - 1. B
        - Question 1 - C
        - 1) D
        """
        answers = {}
        lines = text.strip().split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                continue
            parsed = self._parse_inline_answer(line)
            if parsed:
                q_num, option = parsed
                answers[q_num] = {"correct_option": option, "marks": 1.0}

        if not answers:
            raise ValueError("No answers could be parsed from text content")

        return {"answers": answers}

    def _parse_with_ocr(self, file_path: str, mime_type: str = None) -> Dict:
        """
        Use OCR (Vertex AI Gemini) to extract answer key from PDF or image.
        Sends the file with a specialized prompt asking for structured JSON output.
        """
        if not self.vertex_url:
            raise ValueError("VERTEX_AI_API_KEY not set — cannot use OCR for answer key parsing")

        # Encode file
        with open(file_path, "rb") as f:
            file_data = base64.b64encode(f.read()).decode('utf-8')

        # Detect mime type
        if not mime_type:
            ext = os.path.splitext(file_path)[1].lower()
            mime_map = {
                '.pdf': 'application/pdf',
                '.png': 'image/png',
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
            }
            mime_type = mime_map.get(ext, 'image/png')

        prompt = self._get_answer_key_ocr_prompt()

        payload = {
            "contents": [
                {
                    "role": "user",
                    "parts": [
                        {"text": prompt},
                        {
                            "inline_data": {
                                "mime_type": mime_type,
                                "data": file_data
                            }
                        }
                    ]
                }
            ]
        }

        try:
            headers = self._get_auth_header()
            response = requests.post(
                self.vertex_url,
                headers=headers,
                json=payload,
                timeout=60
            )
            response.raise_for_status()

            result = response.json()
            extracted_text = self._parse_streaming_response(result)
            parsed = self._extract_json(extracted_text)

            if "answers" not in parsed:
                raise ValueError(f"OCR response missing 'answers' field: {extracted_text[:300]}")

            return parsed

        except requests.RequestException as e:
            raise ValueError(f"OCR API request failed: {e}")

    # ──────────────────────────────────────────
    #  Helpers
    # ──────────────────────────────────────────

    def _detect_csv_columns(self, headers: list) -> tuple:
        """Detect question_number, option, and marks columns from headers."""
        if not headers:
            return 0, 1, None  # default: first col = question, second = option

        q_col = None
        opt_col = None
        marks_col = None

        q_keywords = {'question', 'q', 'number', 'q_no', 'qno', 'q.no', 'q no',
                       'question_number', 'question number', 'sl', 'sr', 'sno', "s.no"}
        opt_keywords = {'correct', 'option', 'answer', 'correct_option', 'correct option',
                         'ans', 'key', 'correct_answer', 'correct answer'}
        marks_keywords = {'marks', 'mark', 'score', 'weight', 'points'}

        for idx, header in enumerate(headers):
            h = header.lower().strip()
            if q_col is None and any(kw in h for kw in q_keywords):
                q_col = idx
            elif opt_col is None and any(kw in h for kw in opt_keywords):
                opt_col = idx
            elif marks_col is None and any(kw in h for kw in marks_keywords):
                marks_col = idx

        # Fallback: if we found neither, assume col 0 = question, col 1 = option
        if q_col is None and opt_col is None:
            q_col = 0
            opt_col = 1 if len(headers) > 1 else 0

        # If only one was found, infer the other
        if q_col is not None and opt_col is None:
            opt_col = q_col + 1
        elif opt_col is not None and q_col is None:
            q_col = max(0, opt_col - 1)

        return q_col, opt_col, marks_col

    def _parse_question_number(self, text: str) -> int:
        """Extract integer question number from text like 'Q1', '1.', 'Question 1', '1'"""
        # Remove common prefixes
        cleaned = re.sub(r'^(question|q|no|s\.?no\.?)\s*[.:\-)\s]*', '', text, flags=re.IGNORECASE)
        cleaned = cleaned.strip().rstrip('.):')

        try:
            return int(cleaned)
        except ValueError:
            # Try to find any integer in the text
            match = re.search(r'\d+', text)
            if match:
                return int(match.group())
            raise ValueError(f"Cannot parse question number from: '{text}'")

    def _parse_inline_answer(self, line: str) -> Optional[tuple]:
        """
        Parse a line like:
        'Q1: A', '1. B', 'Question 1 - C', '1) D', '1    A'
        
        Returns: (question_number: int, option: str) or None
        """
        # Pattern: optional prefix (Q/Question) + number + separator + option letter
        patterns = [
            r'(?:question|q)?\.?\s*(\d+)\s*[.:)\-–—\t]+\s*([A-Da-d])\b',
            r'(\d+)\s+([A-Da-d])\s*$',  # Just "1  A"
        ]
        for pattern in patterns:
            match = re.match(pattern, line.strip(), re.IGNORECASE)
            if match:
                q_num = int(match.group(1))
                option = match.group(2).upper()
                return (q_num, option)
        return None

    def _normalize(self, raw: Dict, source_file: str = "") -> AnswerKey:
        """Convert raw parsed dict to AnswerKey model."""
        raw_answers = raw.get("answers", {})
        answers = {}

        for key, value in raw_answers.items():
            q_num = int(key)
            if isinstance(value, dict):
                option = str(value.get("correct_option", "")).strip().upper()
                marks = float(value.get("marks", 1.0))
            elif isinstance(value, str):
                option = value.strip().upper()
                marks = 1.0
            else:
                continue

            # Validate option is a single letter A-D (or allow broader)
            if option and len(option) == 1 and option.isalpha():
                answers[q_num] = AnswerKeyEntry(correct_option=option, marks=marks)
            elif option:
                # Allow non-standard options but warn
                print(f"  ⚠️  Q{q_num}: unusual option '{option}' — keeping as-is")
                answers[q_num] = AnswerKeyEntry(correct_option=option, marks=marks)

        if not answers:
            raise ValueError("No valid answers found after normalization")

        return AnswerKey(
            total_questions=len(answers),
            answers=answers,
            negative_marking=raw.get("negative_marking", 0.0),
            metadata={
                "source_file": source_file,
                "raw_question_count": len(raw_answers),
                "parsed_question_count": len(answers),
            }
        )

    def _get_answer_key_ocr_prompt(self) -> str:
        return """
You are analyzing an ANSWER KEY document for an objective/MCQ examination.

Extract ALL question numbers and their correct options. Return ONLY a valid JSON object (no markdown):

{
    "answers": {
        "1": {"correct_option": "A", "marks": 1},
        "2": {"correct_option": "C", "marks": 1},
        "3": {"correct_option": "B", "marks": 1},
        ...
    },
    "negative_marking": 0
}

Rules:
- Question numbers must be integers (as strings in the JSON keys)
- Options should be single uppercase letters (A, B, C, D)
- If marks per question are visible, include them; otherwise default to 1
- If negative marking info is visible, include it; otherwise set to 0
- Include ALL questions visible in the document
- Return ONLY the JSON object, no explanation text
"""

    def _parse_streaming_response(self, result) -> str:
        """Parse streaming response from Vertex AI."""
        all_text = []
        if isinstance(result, list):
            for chunk in result:
                if "candidates" in chunk:
                    for candidate in chunk["candidates"]:
                        if "content" in candidate and "parts" in candidate["content"]:
                            for part in candidate["content"]["parts"]:
                                if "text" in part:
                                    all_text.append(part["text"])
        elif isinstance(result, dict):
            if "candidates" in result:
                for candidate in result["candidates"]:
                    if "content" in candidate and "parts" in candidate["content"]:
                        for part in candidate["content"]["parts"]:
                            if "text" in part:
                                all_text.append(part["text"])
        return "".join(all_text)

    def _extract_json(self, text: str) -> Dict:
        """Extract JSON from text, handling markdown fences."""
        cleaned = re.sub(r'```json\s*|\s*```', '', text).strip()
        try:
            return json.loads(cleaned)
        except json.JSONDecodeError:
            match = re.search(r'\{.*\}', cleaned, re.DOTALL)
            if match:
                try:
                    return json.loads(match.group())
                except json.JSONDecodeError:
                    pass
            raise ValueError(f"Failed to parse JSON from OCR response: {text[:300]}")
